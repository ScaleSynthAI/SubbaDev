import docx

doc = docx.Document("SubbaTaniparti.docx")

def deep_search(keyword):
    found = []
    # Search paragraphs
    for idx, p in enumerate(doc.paragraphs):
        if keyword.lower() in p.text.lower():
            found.append(f"Paragraph {idx}: {repr(p.text)}")
    # Search tables
    for idx, t in enumerate(doc.tables):
        for r_idx, row in enumerate(t.rows):
            for c_idx, cell in enumerate(row.cells):
                if keyword.lower() in cell.text.lower():
                    found.append(f"Table {idx} R{r_idx}C{c_idx}: {repr(cell.text)}")
    # Search sections
    for idx, sec in enumerate(doc.sections):
        for p_idx, p in enumerate(sec.header.paragraphs):
            if keyword.lower() in p.text.lower():
                found.append(f"Sec {idx} Header P{p_idx}: {repr(p.text)}")
        for p_idx, p in enumerate(sec.footer.paragraphs):
            if keyword.lower() in p.text.lower():
                found.append(f"Sec {idx} Footer P{p_idx}: {repr(p.text)}")
    return found

for word in ["skills", "certifications", "languages", "kubernetes", "cert"]:
    results = deep_search(word)
    print(f"=== Keyword: '{word}' (matches: {len(results)}) ===")
    for r in results[:5]:
        print(f"  {r}")
