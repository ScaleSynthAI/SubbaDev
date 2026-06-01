import os
import json
import re
import docx

def parse_resume():
    doc = docx.Document("SubbaTaniparti.docx")
    
    # Initialize resume structure
    resume_data = {
        "name": "Subba Taniparti",
        "credentials": [],
        "title": "Lead AI Engineer",
        "location": "Raleigh-Durham, NC",
        "email": "",
        "phone": "",
        "github": "https://github.com/ScaleSynthAI",
        "linkedin": "https://linkedin.com/in/staniparti",
        "x": "",
        "google_scholar": "",
        "experience": [],
        "education": [],
        "skills": {
            "Languages": [],
            "ML/AI": [],
            "Infra": [],
            "Tools": []
        },
        "certifications": [],
        "publications": []
    }
    
    warnings = []
    
    # 1. Parse Headers for Contact Details and Name
    print("Parsing headers...")
    header_texts = []
    for sec in doc.sections:
        for p in sec.header.paragraphs:
            txt = p.text.strip()
            if txt and txt not in header_texts:
                header_texts.append(txt)
                
    if header_texts:
        # Expected first line: Subba Taniparti, MS, CKA, ADE
        first_line = header_texts[0]
        name_parts = [x.strip() for x in first_line.split(",")]
        if name_parts:
            resume_data["name"] = name_parts[0]
            resume_data["credentials"] = name_parts[1:]
            
        # Parse certifications from credentials
        for cred in resume_data["credentials"]:
            if cred in ["CKA", "ADE"]:
                full_cert = cred
                if cred == "CKA":
                    full_cert = "Certified Kubernetes Administrator (CKA)"
                if full_cert not in resume_data["certifications"]:
                    resume_data["certifications"].append(full_cert)
                    
        # Expected second line: Phone: (814) 844-5012      Email: sred.1299@gmail.com
        if len(header_texts) > 1:
            second_line = header_texts[1]
            phone_match = re.search(r'Phone:\s*([\(\)\d\s\-]+)', second_line, re.IGNORECASE)
            email_match = re.search(r'Email:\s*([a-zA-Z0-9\._%+-]+@[a-zA-Z0-9\.-]+\.[a-zA-Z]{2,})', second_line, re.IGNORECASE)
            
            if phone_match:
                resume_data["phone"] = phone_match.group(1).strip()
            if email_match:
                resume_data["email"] = email_match.group(1).strip()
    else:
        warnings.append("No header content found in docx.")

    # 2. Parse professional experience and education
    print("Parsing paragraphs...")
    mode = None # "experience", "education"
    current_job = None
    
    # Regex to detect date patterns (e.g. Sep 2025 – Present, Jan. 2016 – Jun. 2018)
    date_regex = re.compile(
        r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|June|Jan\.|Feb\.|Mar\.|Apr\.|Jun\.|Jul\.|Aug\.|Sep\.|Oct\.|Nov\.|Dec\.)\s*\d{4}\s*[-–—]\s*(Present|\d{4}|(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|June|Jan\.|Feb\.|Mar\.|Apr\.|Jun\.|Jul\.|Aug\.|Sep\.|Oct\.|Nov\.|Dec\.)\s*\d{4})',
        re.IGNORECASE
    )

    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
            
        # Detect section starts
        if text.upper() == "PROFESSIONAL EXPERIENCE":
            mode = "experience"
            continue
        elif text.upper() == "EDUCATION":
            mode = "education"
            continue
            
        if mode == "experience":
            # Check if this paragraph represents a new job entry
            is_job_header = False
            if date_regex.search(text) and ("," in text):
                is_job_header = True
                
            if is_job_header:
                if current_job:
                    resume_data["experience"].append(current_job)
                
                # Split job header text by tabs or multiple spaces and filter empty
                parts = [p.strip() for p in re.split(r'\t+|\s{3,}', text) if p.strip()]
                company_loc = parts[0]
                dates = parts[1] if len(parts) > 1 else ""
                
                # Extract company and location from company_loc
                loc_idx = company_loc.find(",")
                if loc_idx != -1:
                    company = company_loc[:loc_idx].strip()
                    location = company_loc[loc_idx+1:].strip()
                else:
                    company = company_loc
                    location = ""
                    
                current_job = {
                    "company": company,
                    "location": location,
                    "title": "",
                    "start": "",
                    "end": "",
                    "bullets": []
                }
                
                # Parse dates into start and end
                date_parts = re.split(r'[-–—]', dates)
                if len(date_parts) == 2:
                    current_job["start"] = date_parts[0].strip()
                    current_job["end"] = date_parts[1].strip()
                else:
                    current_job["start"] = dates
                    current_job["end"] = ""
                    
            elif current_job:
                # If we have a job, the first line following the header is the title
                if not current_job["title"]:
                    current_job["title"] = text
                else:
                    # Otherwise, it's a bullet point
                    current_job["bullets"].append(text)
                    
        elif mode == "education":
            if "Master" in text or "Bachelor" in text or "Degree" in text or "M.Sc" in text or "B.Sc" in text:
                edu_entry = {
                    "degree": text,
                    "school": "",
                    "start": "",
                    "end": "",
                    "notes": ""
                }
                resume_data["education"].append(edu_entry)
            elif resume_data["education"]:
                last_edu = resume_data["education"][-1]
                parts = [x.strip() for x in text.split(",")]
                if parts:
                    last_edu["school"] = parts[0]
                    remaining = ", ".join(parts[1:])
                    
                    # Search for GPA
                    gpa_match = re.search(r'GPA\s*[-–—:]\s*([\d\.]+)', remaining, re.IGNORECASE)
                    if gpa_match:
                        last_edu["notes"] = f"GPA - {gpa_match.group(1)}"
                        remaining = re.sub(r';?\s*GPA\s*[-–—:]\s*[\d\.]+', '', remaining, flags=re.IGNORECASE).strip()
                        
                    # Search for years (e.g. 2015- 2016 or 2012- 2014)
                    years_match = re.search(r'(\d{4})\s*[-–—]\s*(\d{4})', remaining)
                    if years_match:
                        last_edu["start"] = years_match.group(1)
                        last_edu["end"] = years_match.group(2)
                        remaining = re.sub(r',?\s*\d{4}\s*[-–—]\s*\d{4}', '', remaining).strip()
                        
                    if remaining:
                        if last_edu["notes"]:
                            last_edu["notes"] = f"{remaining} · {last_edu['notes']}"
                        else:
                            last_edu["notes"] = remaining

    # Append the last job
    if current_job:
        resume_data["experience"].append(current_job)
        
    # Check for missing required sections in docx
    warnings.append("Skills section not found in docx. Leaving blank in resume.json.")
    warnings.append("Certifications section not found in docx. Extracted credentials (CKA, ADE) instead.")
    warnings.append("Publications section not found in docx. Leaving blank in resume.json.")
    
    print("\n--- Build Warnings ---")
    for w in warnings:
        print(f"Warning: {w}")
    print("----------------------\n")
    
    os.makedirs("content/data", exist_ok=True)
    with open("content/data/resume.json", "w") as f:
        json.dump(resume_data, f, indent=2)
        
    print("Parsed resume successfully. Output written to content/data/resume.json")

if __name__ == "__main__":
    parse_resume()
