import docx

doc = docx.Document("SubbaTaniparti.docx")
print("Total Paragraphs:", len(doc.paragraphs))
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f"[{i}] {p.text[:100]} (style: {p.style.name})")

print("\n--- Tables ---")
for i, t in enumerate(doc.tables):
    print(f"Table {i} rows: {len(t.rows)}, cols: {len(t.columns)}")
    for r_idx, row in enumerate(t.rows[:2]):
        cells = [c.text.strip().replace('\n', ' ') for c in row.cells]
        print(f"  Row {r_idx}: {cells}")
