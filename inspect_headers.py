import docx

doc = docx.Document("SubbaTaniparti.docx")

print("--- Core Properties ---")
props = doc.core_properties
for prop in ['author', 'category', 'comments', 'content_status', 'created', 'identifier', 'keywords', 'language', 'last_modified_by', 'last_printed', 'modified', 'revision', 'subject', 'title', 'version']:
    try:
        val = getattr(props, prop)
        if val:
            print(f"  {prop}: {val}")
    except Exception as e:
        print(f"  {prop}: error {e}")

print("\n--- Sections & Headers/Footers ---")
for idx, sec in enumerate(doc.sections):
    print(f"Section {idx}:")
    header = sec.header
    footer = sec.footer
    
    print("  Header Paragraphs:")
    for p in header.paragraphs:
        print(f"    H: {repr(p.text)}")
    print("  Footer Paragraphs:")
    for p in footer.paragraphs:
        print(f"    F: {repr(p.text)}")
        
    # Check different first page header
    if sec.different_first_page_header_footer:
        print("  Different First Page Header/Footer enabled")
        first_header = sec.first_page_header
        first_footer = sec.first_page_footer
        print("  First Page Header Paragraphs:")
        for p in first_header.paragraphs:
            print(f"    FH: {repr(p.text)}")
        print("  First Page Footer Paragraphs:")
        for p in first_footer.paragraphs:
            print(f"    FF: {repr(p.text)}")
